"""知识库文件解析与清洗模块。

该模块把不同格式的知识文件统一解析成 ``ParsedDocument``：
- md/txt：直接读取文本，按类 Markdown 结构切分。
- pdf：优先读取文本层；没有文本层时使用本地 OCR；表格用 pdfplumber 抽取；
  图片会保存为资产，并在配置多模态模型时生成图片说明。
- docx：读取段落和表格，尽量保留标题/段落/表格的结构信息。
- doc：先通过 LibreOffice 转成 docx，再复用 docx 解析逻辑。

最终输出的 chunk 会进入 RAG 索引，所以这里的目标不是“展示原文排版”，
而是尽量提取对客服回复有用、结构清晰、噪声较少的知识文本。
"""

import csv
import io
import os
import re
import shutil
import subprocess
from hashlib import sha256
from tempfile import TemporaryDirectory
from dataclasses import dataclass
from pathlib import Path

from app.ocr_repair_client import OCRRepairError, repair_ocr_text
from app.vision_client import VisionClientError, describe_image


SUPPORTED_KNOWLEDGE_SUFFIXES = {".md", ".txt", ".pdf", ".docx", ".doc"}
# PDF 中抽取出来的图片会保存到知识库目录下，方便后续在解析报告中查看，
# 也方便未来接入多模态 embedding 或视觉模型时复用。
EXTRACTED_ASSETS_DIR = Path(__file__).resolve().parents[1] / "knowledge_docs" / "extracted_assets"


@dataclass
class ParsedChunk:
    """解析后的最小知识片段。

    page_number 和 section_title 是可选元数据，用于前端展示“知识库依据”时
    告诉用户这个片段来自哪一页、哪个章节。
    """

    content: str
    page_number: int | None = None
    section_title: str = ""


@dataclass
class CleanResult:
    """文本清洗结果，用于生成解析报告。"""

    text: str
    original_chars: int
    cleaned_chars: int
    noise_lines_removed: int


@dataclass
class ParseReport:
    """文件解析报告。

    报告不会参与 RAG 检索本身，但会展示给用户：解析用了什么 parser、
    清洗掉多少噪声、识别出多少页/表格/图片，以及是否存在 OCR 或空页警告。
    """

    file_type: str
    parser: str
    original_chars: int = 0
    cleaned_chars: int = 0
    noise_lines_removed: int = 0
    page_count: int | None = None
    pages_with_text: int | None = None
    section_count: int = 0
    table_count: int = 0
    image_count: int = 0
    ocr_page_count: int = 0
    ocr_llm_repaired_pages: int = 0
    ocr_pages: list[dict] | None = None
    extracted_assets: list[dict] | None = None
    warnings: list[str] | None = None

    def to_dict(self) -> dict:
        return {
            "file_type": self.file_type,
            "parser": self.parser,
            "original_chars": self.original_chars,
            "cleaned_chars": self.cleaned_chars,
            "noise_lines_removed": self.noise_lines_removed,
            "page_count": self.page_count,
            "pages_with_text": self.pages_with_text,
            "section_count": self.section_count,
            "table_count": self.table_count,
            "image_count": self.image_count,
            "ocr_page_count": self.ocr_page_count,
            "ocr_llm_repaired_pages": self.ocr_llm_repaired_pages,
            "ocr_pages": self.ocr_pages or [],
            "extracted_assets": self.extracted_assets or [],
            "warnings": self.warnings or [],
        }


@dataclass
class ParsedDocument:
    """统一的文档解析结果。"""

    title: str
    chunks: list[ParsedChunk]
    report: ParseReport


@dataclass(frozen=True)
class OCRWord:
    """Tesseract TSV 返回的单词及其页面坐标。"""

    text: str
    confidence: float
    block_number: int
    paragraph_number: int
    line_number: int
    left: int
    top: int
    width: int
    height: int


@dataclass(frozen=True)
class OCRTextBlock:
    """由 OCR 单词合并出的文本块，用于恢复多栏阅读顺序。"""

    text: str
    confidence: float
    left: int
    top: int
    width: int
    height: int

    @property
    def right(self) -> int:
        return self.left + self.width

    @property
    def bottom(self) -> int:
        return self.top + self.height

    @property
    def center_x(self) -> float:
        return self.left + self.width / 2


@dataclass(frozen=True)
class OCRPageResult:
    """一页 OCR 的最终文本、质量指标和修复审计信息。"""

    text: str
    raw_text: str = ""
    average_confidence: float = 0.0
    quality_score: float = 0.0
    multi_column: bool = False
    layout_reordered: bool = False
    llm_repaired: bool = False
    repair_model: str = ""
    repair_changes: tuple[str, ...] = ()
    warning: str = ""
    assets: tuple[dict, ...] = ()

    def to_report(self, page_number: int) -> dict:
        return {
            "page_number": page_number,
            "average_confidence": round(self.average_confidence, 4),
            "quality_score": round(self.quality_score, 4),
            "multi_column": self.multi_column,
            "layout_reordered": self.layout_reordered,
            "llm_repaired": self.llm_repaired,
            "repair_model": self.repair_model,
            "repair_changes": list(self.repair_changes),
            "warning": self.warning,
        }


def parse_document(path: Path) -> ParsedDocument:
    """根据文件后缀选择对应解析器。"""
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_KNOWLEDGE_SUFFIXES:
        raise ValueError("Only .md, .txt, .pdf, .docx and .doc knowledge files are supported.")
    if suffix in {".md", ".txt"}:
        return parse_text_document(path)
    if suffix == ".pdf":
        return parse_pdf_document(path)
    if suffix == ".doc":
        return parse_doc_document(path)
    return parse_docx_document(path)


def parse_text_document(path: Path) -> ParsedDocument:
    """解析 md/txt 文本文件。"""
    raw_text = path.read_text(encoding="utf-8-sig")
    clean = clean_text_with_report(raw_text)
    title = extract_title(clean.text, path.stem)
    chunks = split_markdown_like_text(clean.text)
    report = ParseReport(
        file_type=path.suffix.lower().lstrip("."),
        parser="plain-text",
        original_chars=clean.original_chars,
        cleaned_chars=clean.cleaned_chars,
        noise_lines_removed=clean.noise_lines_removed,
        section_count=count_sections(chunks),
    )
    return ParsedDocument(title=title, chunks=chunks, report=report)


def parse_pdf_document(path: Path) -> ParsedDocument:
    """解析 PDF 文件。

    策略是“文本层优先，OCR 兜底”：
    - 如果页面自带可复制文本，直接使用 pypdf 提取，速度快、成本低。
    - 如果某页没有文本层，则调用本地 Tesseract OCR。
    - 表格和图片单独抽取为补充 chunk，避免 PDF 普通文本抽取时丢失结构信息。
    """
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ValueError("PDF parsing dependency is missing. Please install pypdf.") from exc

    reader = PdfReader(str(path))
    page_chunks: list[ParsedChunk] = []
    first_text = ""
    total_original_chars = 0
    total_cleaned_chars = 0
    total_noise_lines = 0
    pages_with_text = 0
    pages_with_ocr = 0
    pages_repaired_by_llm = 0
    ocr_page_reports: list[dict] = []
    ocr_assets: list[dict] = []
    warnings: list[str] = []
    table_chunks, table_count = extract_pdf_table_chunks(path)
    image_chunks, image_assets = extract_pdf_image_reference_chunks(path)

    for page_index, page in enumerate(reader.pages, start=1):
        raw_text = page.extract_text() or ""
        text_source = "text-layer"
        if not raw_text.strip():
            ocr_result = ocr_pdf_page(
                path,
                page_index,
                visual_context=image_context_for_page(image_assets, page_index),
            )
            raw_text = ocr_result.text
            text_source = "ocr"
            if raw_text.strip():
                pages_with_ocr += 1
                pages_repaired_by_llm += int(ocr_result.llm_repaired)
                ocr_page_reports.append(ocr_result.to_report(page_index))
                ocr_assets.extend(ocr_result.assets)
                if ocr_result.warning:
                    warnings.append(f"Page {page_index} OCR warning: {ocr_result.warning}")
            else:
                warnings.append(f"Page {page_index} had no extractable text and OCR returned no text.")

        total_original_chars += len(raw_text)
        clean = clean_text_with_report(raw_text)
        total_cleaned_chars += clean.cleaned_chars
        total_noise_lines += clean.noise_lines_removed
        if not clean.text:
            continue
        pages_with_text += 1
        first_text = first_text or clean.text
        for chunk in split_markdown_like_text(clean.text, page_number=page_index):
            page_chunks.append(chunk)
        if text_source == "ocr":
            warnings.append(f"Page {page_index} was parsed by local OCR.")

    page_chunks.extend(table_chunks)
    page_chunks.extend(image_chunks)

    if not page_chunks:
        raise ValueError("No extractable text found in this PDF. Local OCR also returned no reliable text.")

    if pages_with_text < len(reader.pages):
        warnings.append(f"{len(reader.pages) - pages_with_text} page(s) had no extractable text.")
    parser_name = "pypdf"
    if pages_with_ocr:
        parser_name += " + tesseract-ocr"
    if pages_repaired_by_llm:
        parser_name += " + qwen-ocr-repair"
    report = ParseReport(
        file_type="pdf",
        parser=parser_name,
        original_chars=total_original_chars,
        cleaned_chars=total_cleaned_chars,
        noise_lines_removed=total_noise_lines,
        page_count=len(reader.pages),
        pages_with_text=pages_with_text,
        section_count=count_sections(page_chunks),
        table_count=table_count,
        image_count=len(image_assets),
        ocr_page_count=pages_with_ocr,
        ocr_llm_repaired_pages=pages_repaired_by_llm,
        ocr_pages=ocr_page_reports,
        extracted_assets=[*image_assets, *ocr_assets],
        warnings=warnings,
    )
    return ParsedDocument(title=extract_title(first_text, path.stem), chunks=page_chunks, report=report)


def extract_pdf_table_chunks(path: Path) -> tuple[list[ParsedChunk], int]:
    """用 pdfplumber 提取 PDF 表格并转成 Markdown 表格文本。"""
    try:
        import pdfplumber
    except ImportError:
        return [], 0

    chunks: list[ParsedChunk] = []
    table_count = 0
    try:
        with pdfplumber.open(str(path)) as pdf:
            for page_index, page in enumerate(pdf.pages, start=1):
                tables = page.extract_tables() or []
                for table_index, table in enumerate(tables, start=1):
                    table_text = pdf_table_to_markdown(table)
                    if not table_text:
                        continue
                    table_count += 1
                    title = f"PDF 表格 第 {page_index} 页 表 {table_index}"
                    content = f"## {title}\n\n{table_text}"
                    for chunk in split_markdown_like_text(content, page_number=page_index):
                        chunk.section_title = title
                        chunks.append(chunk)
    except Exception:
        return [], 0
    return chunks, table_count


def pdf_table_to_markdown(table: list[list[str | None]]) -> str:
    """把 pdfplumber 返回的二维表格转换为 Markdown 表格。

    Markdown 表格虽然不是原始版式，但对 RAG 很友好：列关系清晰，
    LLM 也容易理解“表头-单元格”的对应关系。
    """
    rows = [[clean_table_cell(cell) for cell in row] for row in table if row]
    rows = [row for row in rows if any(cell for cell in row)]
    if not rows:
        return ""

    max_columns = max(len(row) for row in rows)
    normalized_rows = [row + [""] * (max_columns - len(row)) for row in rows]
    header = normalized_rows[0]
    body = normalized_rows[1:] or [[""] * max_columns]
    separator = ["---"] * max_columns
    markdown_rows = [header, separator, *body]
    return "\n".join("| " + " | ".join(row) + " |" for row in markdown_rows)


def clean_table_cell(value: str | None) -> str:
    """清理表格单元格，避免竖线破坏 Markdown 表格结构。"""
    text = clean_inline_text(value or "")
    return text.replace("|", "\\|")


def extract_pdf_image_reference_chunks(path: Path) -> tuple[list[ParsedChunk], list[dict]]:
    """抽取 PDF 中的图片资产，并生成可检索的图片引用 chunk。

    图片不一定包含文字，所以这里不会强行把图片当作 OCR 文本处理：
    - 先把原图保存下来，解析报告中可以看到图片路径。
    - 如果配置了多模态模型，则额外生成图片说明，作为 RAG 文本依据。
    - 如果没有多模态能力，也保留页码和图片引用，方便人工回查。
    """
    try:
        import fitz
    except ImportError:
        return [], []

    chunks: list[ParsedChunk] = []
    assets: list[dict] = []
    try:
        document = fitz.open(str(path))
    except Exception:
        return [], []

    try:
        for page_index, page in enumerate(document, start=1):
            images = page.get_images(full=True)
            for image_index, image in enumerate(images, start=1):
                xref = image[0]
                try:
                    image_data = document.extract_image(xref)
                except Exception:
                    continue
                image_bytes = image_data.get("image", b"")
                if not image_bytes:
                    continue
                extension = normalize_image_extension(image_data.get("ext", "png"))
                asset_path = save_extracted_pdf_image(path, page_index, image_index, extension, image_bytes)
                relative_asset = asset_path.relative_to(EXTRACTED_ASSETS_DIR.parent).as_posix()
                image_description, description_error = describe_pdf_image(
                    asset_path,
                    path=path,
                    page_number=page_index,
                    image_index=image_index,
                )
                asset = {
                    "type": "pdf_image",
                    "path": relative_asset,
                    "page_number": page_index,
                    "image_index": image_index,
                    "size_bytes": len(image_bytes),
                    "has_description": bool(image_description),
                    "description_error": description_error,
                }
                if image_description:
                    asset["description"] = image_description
                assets.append(asset)
                title = f"PDF 图片 第 {page_index} 页 图 {image_index}"
                content = (
                    f"## {title}\n\n"
                    f"图片文件: {relative_asset}\n\n"
                    f"{build_pdf_image_chunk_text(image_description)}"
                )
                chunks.append(ParsedChunk(content=content, page_number=page_index, section_title=title))
    finally:
        document.close()

    return chunks, assets


def describe_pdf_image(asset_path: Path, *, path: Path, page_number: int, image_index: int) -> tuple[str, str]:
    context = f"来源 PDF: {path.name}; 页码: {page_number}; 图片序号: {image_index}"
    try:
        description = describe_image(asset_path, context=context) or ""
        if description:
            return description, ""
        return "", "vision_not_configured_or_image_too_large"
    except VisionClientError as exc:
        return "", str(exc)


def build_pdf_image_chunk_text(image_description: str) -> str:
    if image_description:
        return f"图片说明: {image_description}"
    return (
        "图片说明: 该 PDF 页面包含图片，系统已保存原始图片资产。"
        "当前没有可用的多模态图片说明，文本 RAG 只索引此图片引用和页码。"
    )


def save_extracted_pdf_image(path: Path, page_number: int, image_index: int, extension: str, image_bytes: bytes) -> Path:
    EXTRACTED_ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    digest = sha256(image_bytes).hexdigest()[:12]
    safe_stem = safe_asset_name(path.stem)
    asset_name = f"{safe_stem}-p{page_number}-img{image_index}-{digest}.{extension}"
    asset_path = EXTRACTED_ASSETS_DIR / asset_name
    if not asset_path.exists():
        asset_path.write_bytes(image_bytes)
    return asset_path


def safe_asset_name(value: str) -> str:
    safe = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff._-]+", "-", value).strip("-._")
    return safe[:80] or "pdf"


def normalize_image_extension(value: str) -> str:
    extension = re.sub(r"[^0-9A-Za-z]+", "", value.lower()) or "png"
    if extension == "jpeg":
        return "jpg"
    return extension[:8]


def image_context_for_page(image_assets: list[dict], page_number: int) -> str:
    """汇总同页视觉说明，只作为 OCR 校对上下文。"""
    descriptions = [
        str(asset.get("description", "")).strip()
        for asset in image_assets
        if asset.get("page_number") == page_number and asset.get("description")
    ]
    return "\n".join(description for description in descriptions if description)


def ocr_pdf_page(
    path: Path,
    page_number: int,
    *,
    visual_context: str = "",
) -> OCRPageResult:
    """对单页 PDF 做“坐标恢复 + 质量门控 + 可选 LLM 修复”。

    Tesseract 首先输出 TSV，系统根据文本块坐标恢复双栏阅读顺序，并计算平均识别
    置信度和可读质量。只有多栏页或质量低于阈值时才调用 LLM；模型修复失败不会
    中断入库，而是回退到本地排序后的文本。
    """
    pdftoppm = find_executable("pdftoppm")
    tesseract = find_executable("tesseract")
    if not pdftoppm or not tesseract:
        return OCRPageResult(text="", warning="pdftoppm_or_tesseract_missing")

    with TemporaryDirectory(prefix="pdf-ocr-") as tmp:
        output_prefix = Path(tmp) / f"page-{page_number}"
        render_command = [
            str(pdftoppm),
            "-f",
            str(page_number),
            "-l",
            str(page_number),
            "-r",
            "220",
            "-png",
            str(path),
            str(output_prefix),
        ]
        try:
            render_result = subprocess.run(render_command, capture_output=True, text=True, timeout=60)
        except (OSError, subprocess.TimeoutExpired):
            return OCRPageResult(text="", warning="pdf_page_render_failed")
        if render_result.returncode != 0:
            return OCRPageResult(text="", warning="pdf_page_render_failed")

        rendered_pages = sorted(Path(tmp).glob(f"{output_prefix.name}-*.png"))
        if not rendered_pages:
            rendered_pages = sorted(Path(tmp).glob("*.png"))
        if not rendered_pages:
            return OCRPageResult(text="", warning="rendered_page_not_found")

        page_results: list[OCRPageResult] = []
        for image_path in rendered_pages:
            tsv_command = [
                str(tesseract),
                str(image_path),
                "stdout",
                "-l",
                "chi_sim+eng",
                "--psm",
                "3",
                "tsv",
            ]
            try:
                tsv_result = subprocess.run(tsv_command, capture_output=True, text=True, timeout=90)
            except (OSError, subprocess.TimeoutExpired):
                continue
            if tsv_result.returncode != 0 or not tsv_result.stdout.strip():
                continue
            words = parse_tesseract_tsv(tsv_result.stdout)
            if not words:
                continue
            blocks = build_ocr_blocks(words)
            ordered_blocks, multi_column, layout_reordered = order_ocr_blocks(blocks)
            raw_text = "\n\n".join(block.text for block in blocks if block.text.strip()).strip()
            ordered_text = "\n\n".join(block.text for block in ordered_blocks if block.text.strip()).strip()
            average_confidence = weighted_ocr_confidence(words)
            quality_score = score_ocr_quality(ordered_text, average_confidence, ordered_blocks)
            page_results.append(
                finalize_ocr_page(
                    path,
                    page_number,
                    raw_text=raw_text,
                    ordered_text=ordered_text,
                    average_confidence=average_confidence,
                    quality_score=quality_score,
                    multi_column=multi_column,
                    layout_reordered=layout_reordered,
                    visual_context=visual_context,
                )
            )

        if not page_results:
            return OCRPageResult(text="", warning="tesseract_returned_no_text")
        if len(page_results) == 1:
            return page_results[0]

        combined_text = "\n\n".join(result.text for result in page_results if result.text.strip())
        combined_raw = "\n\n".join(result.raw_text for result in page_results if result.raw_text.strip())
        all_assets = tuple(asset for result in page_results for asset in result.assets)
        return OCRPageResult(
            text=combined_text,
            raw_text=combined_raw,
            average_confidence=sum(result.average_confidence for result in page_results) / len(page_results),
            quality_score=sum(result.quality_score for result in page_results) / len(page_results),
            multi_column=any(result.multi_column for result in page_results),
            layout_reordered=any(result.layout_reordered for result in page_results),
            llm_repaired=any(result.llm_repaired for result in page_results),
            repair_model=next((result.repair_model for result in page_results if result.repair_model), ""),
            repair_changes=tuple(change for result in page_results for change in result.repair_changes),
            warning="; ".join(result.warning for result in page_results if result.warning),
            assets=all_assets,
        )


def parse_tesseract_tsv(tsv_text: str) -> list[OCRWord]:
    """把 Tesseract TSV 转成带坐标的单词，忽略空白和无效置信度记录。"""
    words: list[OCRWord] = []
    reader = csv.DictReader(io.StringIO(tsv_text), delimiter="\t")
    for row in reader:
        text = clean_inline_text(row.get("text") or "")
        if not text:
            continue
        try:
            confidence = float(row.get("conf") or -1)
            if confidence < 0:
                continue
            words.append(
                OCRWord(
                    text=text,
                    confidence=confidence / 100,
                    block_number=int(row.get("block_num") or 0),
                    paragraph_number=int(row.get("par_num") or 0),
                    line_number=int(row.get("line_num") or 0),
                    left=int(row.get("left") or 0),
                    top=int(row.get("top") or 0),
                    width=int(row.get("width") or 0),
                    height=int(row.get("height") or 0),
                )
            )
        except (TypeError, ValueError):
            continue
    return words


def build_ocr_blocks(words: list[OCRWord]) -> list[OCRTextBlock]:
    """按 Tesseract block/paragraph/line 合并单词，同时保留块坐标。"""
    grouped: dict[tuple[int, int], list[OCRWord]] = {}
    for word in words:
        grouped.setdefault((word.block_number, word.paragraph_number), []).append(word)

    blocks: list[OCRTextBlock] = []
    for block_words in grouped.values():
        line_groups: dict[int, list[OCRWord]] = {}
        for word in block_words:
            line_groups.setdefault(word.line_number, []).append(word)
        line_texts: list[tuple[int, str]] = []
        for line_words in line_groups.values():
            ordered_words = sorted(line_words, key=lambda item: item.left)
            line_text = smart_join_ocr_words([item.text for item in ordered_words])
            line_texts.append((min(item.top for item in ordered_words), line_text))
        block_text = "\n".join(text for _, text in sorted(line_texts) if text).strip()
        if not block_text:
            continue
        left = min(word.left for word in block_words)
        top = min(word.top for word in block_words)
        right = max(word.left + word.width for word in block_words)
        bottom = max(word.top + word.height for word in block_words)
        total_area = sum(max(1, word.width * word.height) for word in block_words)
        confidence = safe_ratio(
            sum(word.confidence * max(1, word.width * word.height) for word in block_words),
            total_area,
        )
        blocks.append(
            OCRTextBlock(
                text=block_text,
                confidence=confidence,
                left=left,
                top=top,
                width=max(1, right - left),
                height=max(1, bottom - top),
            )
        )
    return sorted(blocks, key=lambda item: (item.top, item.left))


def order_ocr_blocks(blocks: list[OCRTextBlock]) -> tuple[list[OCRTextBlock], bool, bool]:
    """识别左右双栏并按“左栏读完，再读右栏”恢复阅读顺序。"""
    if len(blocks) < 4:
        return blocks, False, False
    page_left = min(block.left for block in blocks)
    page_right = max(block.right for block in blocks)
    page_width = max(1, page_right - page_left)
    center = page_left + page_width / 2
    narrow_blocks = [block for block in blocks if block.width <= page_width * 0.58]
    left_column = [
        block for block in narrow_blocks
        if block.center_x < center - page_width * 0.02
        and block.right <= center + page_width * 0.08
    ]
    right_column = [
        block for block in narrow_blocks
        if block.center_x > center + page_width * 0.02
        and block.left >= center - page_width * 0.08
    ]
    if len(left_column) < 2 or len(right_column) < 2:
        return blocks, False, False

    overlap_top = max(min(block.top for block in left_column), min(block.top for block in right_column))
    overlap_bottom = min(max(block.bottom for block in left_column), max(block.bottom for block in right_column))
    if overlap_bottom <= overlap_top:
        return blocks, False, False

    column_top = min(min(block.top for block in left_column), min(block.top for block in right_column))
    column_bottom = max(max(block.bottom for block in left_column), max(block.bottom for block in right_column))
    left_ids = {id(block) for block in left_column}
    right_ids = {id(block) for block in right_column}
    headers = [block for block in blocks if block.bottom < column_top and id(block) not in left_ids | right_ids]
    footers = [block for block in blocks if block.top > column_bottom and id(block) not in left_ids | right_ids]
    middle_other = [
        block
        for block in blocks
        if id(block) not in left_ids | right_ids
        and block not in headers
        and block not in footers
    ]
    ordered = [
        *sorted(headers, key=lambda item: (item.top, item.left)),
        *sorted(left_column, key=lambda item: (item.top, item.left)),
        *sorted(middle_other, key=lambda item: (item.top, item.left)),
        *sorted(right_column, key=lambda item: (item.top, item.left)),
        *sorted(footers, key=lambda item: (item.top, item.left)),
    ]
    original_ids = [id(block) for block in blocks]
    return ordered, True, [id(block) for block in ordered] != original_ids


def smart_join_ocr_words(words: list[str]) -> str:
    """连接 OCR 单词；中文字符间不加空格，英文单词间保留空格。"""
    result = ""
    for word in words:
        if not result:
            result = word
            continue
        previous = result[-1]
        current = word[0]
        needs_space = (
            previous.isascii()
            and current.isascii()
            and previous.isalnum()
            and current.isalnum()
        )
        result += (" " if needs_space else "") + word
    return result.strip()


def weighted_ocr_confidence(words: list[OCRWord]) -> float:
    weights = [max(1, len(word.text)) for word in words]
    return safe_ratio(
        sum(word.confidence * weight for word, weight in zip(words, weights)),
        sum(weights),
    )


def safe_ratio(numerator: float, denominator: float) -> float:
    return numerator / denominator if denominator else 0.0


def score_ocr_quality(
    text: str,
    average_confidence: float,
    blocks: list[OCRTextBlock],
) -> float:
    """综合字符可读性、Tesseract 置信度和碎片程度计算 0～1 质量分。"""
    if not text.strip():
        return 0.0
    readability = readable_ratio(text)
    suspicious_chars = len(re.findall(r"[�□■◆◇]{1}|\\x[0-9a-fA-F]{2}", text))
    symbol_penalty = min(1.0, suspicious_chars / max(1, len(text)) * 20)
    short_blocks = sum(len(block.text.strip()) < 8 for block in blocks)
    structure_score = 1 - safe_ratio(short_blocks, max(1, len(blocks)))
    quality = (
        average_confidence * 0.55
        + readability * 0.30
        + structure_score * 0.15
        - symbol_penalty * 0.20
    )
    return round(max(0.0, min(1.0, quality)), 4)


def finalize_ocr_page(
    path: Path,
    page_number: int,
    *,
    raw_text: str,
    ordered_text: str,
    average_confidence: float,
    quality_score: float,
    multi_column: bool,
    layout_reordered: bool,
    visual_context: str = "",
) -> OCRPageResult:
    """按质量门控决定是否调用 LLM，并保存原始与最终文本供审计。"""
    final_text = ordered_text or raw_text
    repair_model = ""
    repair_changes: tuple[str, ...] = ()
    warning = ""
    # Tesseract 可能对“形似但语义错误”的汉字给出很高字符置信度，例如把“策略”
    # 识别成“梨略”。默认阈值因此保持得较严格，让这类页面也进入保守的语义校对。
    threshold = float(os.getenv("OCR_LLM_REPAIR_THRESHOLD", "0.96"))
    should_repair = bool(final_text) and (quality_score < threshold or multi_column)
    if should_repair:
        try:
            repaired = repair_ocr_text(
                final_text,
                document_name=path.name,
                page_number=page_number,
                quality_score=quality_score,
                visual_context=visual_context,
            )
            if repaired:
                final_text = repaired.text
                repair_model = repaired.model
                repair_changes = repaired.changes
        except OCRRepairError as exc:
            warning = str(exc)

    raw_asset = save_ocr_text_asset(path, page_number, "raw", raw_text)
    final_asset = save_ocr_text_asset(path, page_number, "repaired" if repair_model else "ordered", final_text)
    assets = (
        build_ocr_asset(raw_asset, page_number, "ocr_raw"),
        build_ocr_asset(final_asset, page_number, "ocr_repaired" if repair_model else "ocr_ordered"),
    )
    return OCRPageResult(
        text=final_text,
        raw_text=raw_text,
        average_confidence=average_confidence,
        quality_score=quality_score,
        multi_column=multi_column,
        layout_reordered=layout_reordered,
        llm_repaired=bool(repair_model),
        repair_model=repair_model,
        repair_changes=repair_changes,
        warning=warning,
        assets=assets,
    )


def save_ocr_text_asset(path: Path, page_number: int, variant: str, text: str) -> Path:
    EXTRACTED_ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    digest = sha256(text.encode("utf-8")).hexdigest()[:12]
    asset_name = f"{safe_asset_name(path.stem)}-p{page_number}-ocr-{variant}-{digest}.txt"
    asset_path = EXTRACTED_ASSETS_DIR / asset_name
    if not asset_path.exists():
        asset_path.write_text(text, encoding="utf-8")
    return asset_path


def build_ocr_asset(path: Path, page_number: int, asset_type: str) -> dict:
    return {
        "type": asset_type,
        "path": path.relative_to(EXTRACTED_ASSETS_DIR.parent).as_posix(),
        "page_number": page_number,
        "size_bytes": path.stat().st_size,
    }


def parse_docx_document(path: Path) -> ParsedDocument:
    """解析 DOCX 文档。

    段落会按 heading 样式保留章节结构；表格会转成文本块。
    这样切分时可以优先按章节和段落切，而不是简单按固定字符数硬切。
    """
    try:
        from docx import Document
    except ImportError as exc:
        raise ValueError("DOCX parsing dependency is missing. Please install python-docx.") from exc

    document = Document(str(path))
    blocks: list[str] = []
    current_section = ""
    raw_chars = 0

    for paragraph in document.paragraphs:
        raw_chars += len(paragraph.text or "")
        text = clean_inline_text(paragraph.text)
        if not text:
            continue
        style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
        if style_name.startswith("heading"):
            current_section = text
            blocks.append(f"## {text}")
        elif current_section:
            blocks.append(text)
        else:
            blocks.append(text)

    for table in document.tables:
        table_text = table_to_text(table)
        raw_chars += len(table_text)
        if table_text:
            blocks.append(table_text)

    clean = clean_text_with_report("\n\n".join(blocks))
    if not clean.text:
        raise ValueError("No extractable text found in this DOCX.")
    chunks = split_markdown_like_text(clean.text)
    report = ParseReport(
        file_type="docx",
        parser="python-docx",
        original_chars=raw_chars,
        cleaned_chars=clean.cleaned_chars,
        noise_lines_removed=clean.noise_lines_removed,
        section_count=count_sections(chunks),
        table_count=len(document.tables),
    )
    return ParsedDocument(title=extract_title(clean.text, path.stem), chunks=chunks, report=report)


def parse_doc_document(path: Path) -> ParsedDocument:
    """解析 legacy .doc 文件。

    优先用 LibreOffice 转成 docx，因为这种方式能最大程度保留结构。
    如果转换失败，再尝试从 OLE 二进制流中提取可读字符串作为兜底。
    """
    conversion_error = ""
    try:
        converted_path = convert_doc_to_docx(path)
        try:
            parsed = parse_docx_document(converted_path)
            parsed.report.file_type = "doc"
            parsed.report.parser = "LibreOffice -> python-docx"
            parsed.report.warnings = [*(parsed.report.warnings or []), "Converted from legacy .doc to .docx before parsing."]
            return parsed
        finally:
            converted_path.unlink(missing_ok=True)
    except ValueError as exc:
        conversion_error = str(exc)

    try:
        import olefile
    except ImportError as exc:
        raise ValueError("DOC parsing dependency is missing. Please install olefile.") from exc

    if not olefile.isOleFile(str(path)):
        raise ValueError(f"{conversion_error} Fallback failed: this .doc file is not a valid legacy Word OLE document.")

    text_parts: list[str] = []
    with olefile.OleFileIO(str(path)) as ole:
        stream_names = ["/".join(item) for item in ole.listdir(streams=True)]
        preferred_streams = [name for name in stream_names if name in {"WordDocument", "1Table", "0Table"}]
        candidate_streams = preferred_streams or stream_names

        for stream_name in candidate_streams:
            try:
                data = ole.openstream(stream_name).read()
            except OSError:
                continue
            text_parts.extend(extract_readable_strings(data))

    clean = clean_text_with_report("\n\n".join(text_parts))
    if len(clean.text) < 20:
        raise ValueError(
            f"{conversion_error} Fallback failed: no reliable text could be extracted from this .doc file. "
            "Please convert it to .docx or PDF and upload again."
        )
    chunks = split_markdown_like_text(clean.text)
    report = ParseReport(
        file_type="doc",
        parser="olefile fallback",
        original_chars=sum(len(part) for part in text_parts),
        cleaned_chars=clean.cleaned_chars,
        noise_lines_removed=clean.noise_lines_removed,
        section_count=count_sections(chunks),
        warnings=[conversion_error, "Used OLE text extraction fallback; formatting may be incomplete."],
    )
    return ParsedDocument(title=extract_title(clean.text, path.stem), chunks=chunks, report=report)


def convert_doc_to_docx(path: Path) -> Path:
    """调用 LibreOffice/soffice 把 .doc 转换成 .docx。"""
    soffice = find_soffice()
    if not soffice:
        raise ValueError("LibreOffice was not found, so .doc could not be converted to .docx.")

    with TemporaryDirectory() as tmp:
        output_dir = Path(tmp)
        command = [
            str(soffice),
            "--headless",
            "--convert-to",
            "docx",
            "--outdir",
            str(output_dir),
            str(path),
        ]
        result = subprocess.run(command, capture_output=True, text=True, timeout=60)
        converted = output_dir / f"{path.stem}.docx"
        if result.returncode != 0 or not converted.exists():
            message = (result.stderr or result.stdout or "LibreOffice did not produce a DOCX file.").strip()
            raise ValueError(f"LibreOffice conversion failed: {message}")

        stable_path = path.with_suffix(".converted.docx")
        stable_path.write_bytes(converted.read_bytes())
        return stable_path


def find_soffice() -> Path | None:
    """查找 LibreOffice 可执行文件。

    Windows 下用户可能安装在 C 盘默认目录，也可能按项目要求安装到 A/B 盘。
    因此这里先查 PATH，再查几个常见安装路径。
    """
    found = find_executable("soffice")
    if found:
        return found

    candidates = [
        Path("C:/Program Files/LibreOffice/program/soffice.exe"),
        Path("C:/Program Files (x86)/LibreOffice/program/soffice.exe"),
        Path("B:/LibreOffice/program/soffice.exe"),
        Path("B:/tools/LibreOffice/program/soffice.exe"),
        Path("A:/LibreOffice/program/soffice.exe"),
        Path("A:/tools/LibreOffice/program/soffice.exe"),
    ]
    return next((candidate for candidate in candidates if candidate.exists()), None)


def find_executable(name: str) -> Path | None:
    """从 PATH 中查找命令行工具。"""
    candidates = [name]
    if not name.endswith(".exe"):
        candidates.append(f"{name}.exe")
    for candidate in candidates:
        found = shutil.which(candidate)
        if found:
            return Path(found)
    return None


def clean_text(text: str) -> str:
    """只返回清洗后的文本，供不需要报告的旧路径使用。"""
    return clean_text_with_report(text).text


def clean_text_with_report(text: str) -> CleanResult:
    """清洗文档文本并统计清洗报告。

    清洗会去掉控制字符、明显乱码行、重复空白和低可读内容。
    这里保留统计信息，是为了前端能解释“为什么这个文档被切成这些 chunk”。
    """
    original_chars = len(text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [clean_inline_text(line) for line in text.split("\n")]
    noise_lines_removed = sum(1 for line in lines if is_noise_line(line))
    lines = [line for line in lines if not is_noise_line(line)]
    text = "\n".join(lines)
    text = re.sub(r"(?<![。！？.!?:：；;])\n(?!\n|[#\-*0-9])", " ", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = text.strip()
    return CleanResult(
        text=text,
        original_chars=original_chars,
        cleaned_chars=len(text),
        noise_lines_removed=noise_lines_removed,
    )


def extract_readable_strings(data: bytes) -> list[str]:
    """从 legacy .doc 的二进制流中提取可读文本片段。

    这是 LibreOffice 转换失败后的兜底路径，不能保证完整排版，
    但可以尽量避免旧 Word 文件完全无法入库。
    """
    candidates: list[str] = []
    candidates.extend(extract_decoded_runs(data.decode("utf-16le", errors="ignore")))
    candidates.extend(extract_decoded_runs(data.decode("latin-1", errors="ignore")))

    seen: set[str] = set()
    readable: list[str] = []
    for candidate in candidates:
        normalized = clean_inline_text(candidate)
        if len(normalized) < 4 or normalized in seen:
            continue
        if readable_ratio(normalized) < 0.55:
            continue
        seen.add(normalized)
        readable.append(normalized)
    return readable


def extract_decoded_runs(text: str) -> list[str]:
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]+", "\n", text)
    pattern = r"[\u4e00-\u9fffA-Za-z0-9，。！？、；：,.!?;:'\"()\[\]《》<>/\-_\s]{4,}"
    return [match.group(0).strip() for match in re.finditer(pattern, text) if match.group(0).strip()]


def readable_ratio(text: str) -> float:
    """计算一行文本中可读字符占比，用于过滤明显乱码。"""
    if not text:
        return 0.0
    readable_chars = re.findall(r"[\u4e00-\u9fffA-Za-z0-9，。！？、；：,.!?;:'\"()\[\]《》<>/\-_\s]", text)
    return len(readable_chars) / len(text)


def clean_inline_text(text: str) -> str:
    """清理单行文本中的连续空白。"""
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]+", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def is_noise_line(line: str) -> bool:
    """判断一行是否像噪声。

    过短、可读比例过低、乱码符号过多的行会被过滤，减少 chunk 中的无效内容。
    """
    if not line:
        return False
    if re.fullmatch(r"[-_—=]{3,}", line):
        return True
    if re.fullmatch(r"(page\s*)?\d+(/\d+)?", line, flags=re.IGNORECASE):
        return True
    if re.fullmatch(r"第\s*\d+\s*页\s*(共\s*\d+\s*页)?", line):
        return True
    return False


def split_markdown_like_text(text: str, page_number: int | None = None, max_chars: int = 900, overlap_chars: int = 100) -> list[ParsedChunk]:
    """结构感知切分入口。

    切分优先级是：
    1. 先识别章节标题，把文档拆成章节。
    2. 章节内优先保留段落和表格块。
    3. 如果块太长，再按行、句子递归切。
    4. 仍然超长时才按固定字符数兜底。

    ``overlap_chars`` 用来在相邻 chunk 之间保留少量尾部上下文，降低句子被切断后
    检索缺少上下文的问题。
    """
    sections = split_into_sections(text)
    chunks: list[ParsedChunk] = []

    for section_title, section_blocks in sections:
        chunks.extend(
            split_blocks_recursively(
                section_blocks,
                page_number=page_number,
                section_title=section_title,
                max_chars=max_chars,
                overlap_chars=overlap_chars,
            )
        )
    return chunks


def split_into_sections(text: str) -> list[tuple[str, list[str]]]:
    """根据 Markdown 标题或独立标题块拆分章节。"""
    blocks = [block.strip() for block in re.split(r"\n\s*\n", text) if block.strip()]
    sections: list[tuple[str, list[str]]] = []
    current_title = ""
    current_blocks: list[str] = []

    for block in blocks:
        heading = extract_section_heading(block)
        if heading and current_blocks:
            sections.append((current_title, current_blocks))
            current_blocks = []
        if heading:
            current_title = heading
            if is_standalone_heading_block(block):
                continue
        current_blocks.append(block)

    if current_blocks:
        sections.append((current_title, current_blocks))
    return sections


def split_blocks_recursively(
    blocks: list[str],
    page_number: int | None,
    section_title: str,
    max_chars: int,
    overlap_chars: int,
) -> list[ParsedChunk]:
    """对一个章节内的 block 做递归切分并打包成 chunk。"""
    pieces: list[str] = []
    for block in blocks:
        pieces.extend(split_oversized_block(block, max_chars=max_chars))
    return pack_chunk_pieces(
        pieces,
        page_number=page_number,
        section_title=section_title,
        max_chars=max_chars,
        overlap_chars=overlap_chars,
    )


def split_oversized_block(block: str, max_chars: int) -> list[str]:
    """切分超长 block。

    这里的顺序体现了“尽量不破坏语义结构”的原则：
    多行文本先按行切；单行长文本再按句子切；句子仍然过长才按字符硬切。
    """
    block = block.strip()
    if not block:
        return []
    if len(block) <= max_chars:
        return [block]

    lines = [line.strip() for line in block.splitlines() if line.strip()]
    if len(lines) > 1:
        line_pieces: list[str] = []
        for line in lines:
            line_pieces.extend(split_oversized_block(line, max_chars=max_chars))
        return line_pieces

    sentences = split_sentences(block)
    if len(sentences) > 1:
        return pack_text_units(sentences, max_chars=max_chars)

    return split_by_chars(block, max_chars=max_chars)


def split_sentences(text: str) -> list[str]:
    sentence_pattern = r"[^。！？!?；;.\n]+[。！？!?；;.]?"
    sentences = [match.group(0).strip() for match in re.finditer(sentence_pattern, text) if match.group(0).strip()]
    return sentences or [text.strip()]


def pack_text_units(units: list[str], max_chars: int) -> list[str]:
    """把句子/短文本单元重新打包到 max_chars 附近。"""
    packed: list[str] = []
    current = ""
    for unit in units:
        if len(unit) > max_chars:
            if current:
                packed.append(current.strip())
                current = ""
            packed.extend(split_by_chars(unit, max_chars=max_chars))
            continue
        candidate = join_sentence_units(current, unit) if current else unit
        if current and len(candidate) > max_chars:
            packed.append(current.strip())
            current = unit
        else:
            current = candidate
    if current:
        packed.append(current.strip())
    return packed


def join_sentence_units(left: str, right: str) -> str:
    if not left:
        return right
    separator = "" if left[-1] in "。！？；" else " "
    return f"{left}{separator}{right}".strip()


def split_by_chars(text: str, max_chars: int) -> list[str]:
    """最后兜底的固定字符切分。"""
    return [text[index : index + max_chars].strip() for index in range(0, len(text), max_chars) if text[index : index + max_chars].strip()]


def pack_chunk_pieces(
    pieces: list[str],
    page_number: int | None,
    section_title: str,
    max_chars: int,
    overlap_chars: int,
) -> list[ParsedChunk]:
    """把切分后的 pieces 组装成最终 ParsedChunk。

    组装时会尽量把短段落合并，避免 chunk 过碎；当长度超限时再开新 chunk，
    并把上一个 chunk 的尾部作为 overlap 拼到新 chunk 前面。
    """
    chunks: list[ParsedChunk] = []
    current = ""

    for piece in pieces:
        separator = "\n\n" if "\n" in piece or piece.startswith("#") else " "
        candidate = f"{current}{separator}{piece}".strip() if current else piece
        if current and len(candidate) > max_chars:
            chunks.append(ParsedChunk(content=current.strip(), page_number=page_number, section_title=section_title))
            overlap = tail_overlap(current, overlap_chars)
            current = f"{overlap}{separator}{piece}".strip() if overlap else piece
            if len(current) > max_chars:
                current = piece
                if len(current) > max_chars:
                    for part in split_by_chars(current, max_chars=max_chars):
                        chunks.append(ParsedChunk(content=part.strip(), page_number=page_number, section_title=section_title))
                    current = ""
        else:
            current = candidate

    if current:
        chunks.append(ParsedChunk(content=current.strip(), page_number=page_number, section_title=section_title))
    return chunks


def count_sections(chunks: list[ParsedChunk]) -> int:
    """统计解析结果中识别到的章节数量。"""
    return len({chunk.section_title for chunk in chunks if chunk.section_title})


def tail_overlap(text: str, overlap_chars: int) -> str:
    """截取 chunk 尾部重叠文本，优先从句子边界开始。"""
    if overlap_chars <= 0 or len(text) <= overlap_chars:
        return ""
    tail = text[-overlap_chars:].strip()
    sentence_start = max(tail.rfind("。"), tail.rfind("."), tail.rfind("\n"))
    return tail[sentence_start + 1 :].strip() if sentence_start >= 0 else tail


def extract_title(text: str, fallback: str) -> str:
    for line in text.splitlines():
        stripped = line.strip("# ").strip()
        if stripped and not is_noise_line(stripped):
            if len(stripped) > 80:
                break
            return stripped[:120]
    return fallback.replace("_", " ").replace("-", " ").title()


def extract_heading(block: str) -> str:
    first_line = block.splitlines()[0].strip()
    markdown_heading = re.match(r"^#{1,4}\s+(.+)$", first_line)
    if markdown_heading:
        return markdown_heading.group(1).strip()[:120]
    if len(first_line) <= 60 and re.match(r"^(\d+(\.\d+)*[、. ]*)?[\u4e00-\u9fffA-Za-z].*", first_line):
        if len(block.splitlines()) == 1 or first_line.endswith(("：", ":")):
            return first_line.strip("：:")[:120]
    return ""


def extract_section_heading(block: str) -> str:
    first_line = block.splitlines()[0].strip()
    markdown_heading = re.match(r"^#{1,4}\s+(.+)$", first_line)
    if markdown_heading:
        return markdown_heading.group(1).strip()[:120]
    if len(block.splitlines()) != 1 or len(first_line) > 80:
        return ""
    numbered_heading = re.match(r"^\d+(\.\d+)*[.、 ]+[\u4e00-\u9fffA-Za-z].+", first_line)
    if numbered_heading:
        return first_line.strip(":：")[:120]
    if first_line.endswith((":","：")):
        return first_line.strip(":：")[:120]
    return ""


def is_standalone_heading_block(block: str) -> bool:
    first_line = block.splitlines()[0].strip() if block.splitlines() else ""
    return len(block.splitlines()) == 1 and bool(re.match(r"^#{1,4}\s+.+$", first_line))


def table_to_text(table) -> str:
    rows: list[str] = []
    for row in table.rows:
        cells = [clean_inline_text(cell.text) for cell in row.cells]
        cells = [cell for cell in cells if cell]
        if cells:
            rows.append(" | ".join(cells))
    return "\n".join(rows)
